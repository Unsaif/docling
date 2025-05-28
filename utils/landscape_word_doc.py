import pandas as pd
import pathlib
from docx import Document
from docx.shared import Pt
from docx.enum.section import WD_ORIENT

def row_to_landscape_doc(df: pd.DataFrame, out_path: str|pathlib.Path) -> None:
    """Transpose df, write one landscape Word file at 7.5 pt."""
    df_t = df.T.reset_index(names=["Field"])

    # 3️⃣  create the Word document in landscape
    doc = Document()
    sec = doc.sections[-1]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width

    # 4️⃣  build a table with **no header row**
    rows, cols = df_t.shape
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"

    # 5️⃣  dump every cell directly
    for i, row in df_t.iterrows():
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = str(val)

    # 6️⃣  set 7.5-pt font everywhere
    for cell in table._cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(7.5)

    doc.save(out_path)